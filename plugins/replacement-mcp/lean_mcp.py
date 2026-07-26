# /// script
# dependencies = [
#   "fastmcp",
#   "httpx",
#   "trafilatura",
#   "ddgs",
#   "openpyxl",
#   "python-docx",
#   "pymupdf",
#   "pyyaml"
# ]
# ///

import os
import re
import json
import csv
import io
import subprocess
from collections import deque
from pathlib import Path
from typing import Literal, Optional, Any, Dict, List, Union
from urllib.parse import urlparse
import yaml
import fitz  # PyMuPDF
from fastmcp import FastMCP

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# ---------------------------------------------------------------------------
# Robust YAML config loading
# ---------------------------------------------------------------------------
CONFIG_PATH = Path(__file__).resolve().parent / "tool_prompts.yaml"
if CONFIG_PATH.exists():
    with open(CONFIG_PATH, "r", encoding="utf-8") as f:
        CONFIG = yaml.safe_load(f)
else:
    CONFIG = {}

PROMPTS = CONFIG
THRESH = PROMPTS.get("thresholds", {})

# ---------------------------------------------------------------------------
# Standardized JSON Response Helpers & Dynamic Error Recovery Hints (Item 5)
# ---------------------------------------------------------------------------
def success_response(
    data: Any,
    message: Optional[str] = None,
    truncated: bool = False,
    metadata: Optional[Dict[str, Any]] = None,
) -> str:
    """Returns a standardized JSON success response."""
    payload: Dict[str, Any] = {
        "status": "success",
        "truncated": truncated,
        "data": data,
    }
    if message:
        payload["message"] = message
    if metadata:
        payload["metadata"] = metadata
    return json.dumps(payload, indent=2, ensure_ascii=False)


def error_response(
    code: str,
    message: str,
    detail: Optional[str] = None,
    hint: Optional[str] = None,
) -> str:
    """Returns a standardized JSON error response with automated recovery hints."""
    payload: Dict[str, Any] = {
        "status": "error",
        "error_code": code,
        "message": message,
    }
    if detail:
        payload["detail"] = detail

    # Item 5: Automatic recovery hints to help small models self-correct
    if not hint:
        if "NOT_FOUND" in code or "MISSING" in code:
            hint = "Verify path spelling or call lean_analyze_workspace to check available files."
        elif "CORRUPT" in code or "PARSE" in code:
            hint = "File may be damaged or password-protected. Verify format."
        elif "BAD_RANGE" in code or "OUT_OF_BOUNDS" in code:
            hint = "Inspect dimensions or line counts before specifying bounds."
        elif "TARGET_NOT_FOUND" in code:
            hint = "Use lean_file_read first to confirm exact string formatting or line numbers."
        elif "SEARCH" in code or "SCRAPE" in code:
            hint = "Broaden search keywords or check domain connectivity."

    if hint:
        payload["hint"] = hint
    return json.dumps(payload, indent=2, ensure_ascii=False)


# ---------------------------------------------------------------------------
# In-Tool Keyword RAG Helper (Item 3)
# ---------------------------------------------------------------------------
def _filter_by_query(
    items: List[str], query: Optional[str] = None, max_results: int = 50
) -> tuple[List[str], bool]:
    """Filters lines or paragraphs by keyword match score."""
    if not query or not query.strip():
        return items[:max_results], len(items) > max_results

    query_words = set(re.findall(r"\w+", query.lower()))
    if not query_words:
        return items[:max_results], len(items) > max_results

    scored = []
    for idx, item in enumerate(items):
        item_words = set(re.findall(r"\w+", item.lower()))
        score = len(query_words.intersection(item_words))
        if score > 0:
            scored.append((score, idx, item))

    if not scored:
        fallback = [f"[No direct matches for query '{query}'. Showing top section]"] + items[:max_results]
        return fallback, len(items) > max_results

    scored.sort(key=lambda x: x[0], reverse=True)
    results = [item for _, _, item in scored[:max_results]]
    return results, len(scored) > max_results


# ---------------------------------------------------------------------------
# Initialize MCP Server
# ---------------------------------------------------------------------------
mcp = FastMCP("lean-goose-mcp", instructions=PROMPTS.get("server_instructions", ""))

# ---------------------------------------------------------------------------
# Persistent directories
# ---------------------------------------------------------------------------
CACHE_DIR = Path.home() / ".cache" / "lean-goose-mcp"
SCRATCH_FILE = CACHE_DIR / "scratchpad.json"
CACHE_DIR.mkdir(parents=True, exist_ok=True)


# ---------------------------------------------------------------------------
# Helper functions
# ---------------------------------------------------------------------------
def _strip_ansi(text: str) -> str:
    return re.sub(r"\x1B(?:[@-Z\\-_]|\[[0-?]*[ -/]*[@-~])", "", text)


def _load_scratchpad() -> dict[str, str]:
    if SCRATCH_FILE.exists():
        with open(SCRATCH_FILE, "r", encoding="utf-8") as f:
            return json.load(f)
    return {}


def _save_scratchpad(data: dict[str, str]) -> None:
    SCRATCH_FILE.parent.mkdir(parents=True, exist_ok=True)
    with open(SCRATCH_FILE, "w", encoding="utf-8") as f:
        json.dump(data, f, indent=2)


# ===========================================================================
# System Tools: shell & analyze_workspace
# ===========================================================================
@mcp.tool(name="lean_shell")
def shell(command: str, dry_run: bool = False) -> str:
    """Runs a shell command and returns truncated stdout/stderr. Set dry_run=True to preview without executing."""
    if dry_run:
        return success_response({"command": command}, message="[DRY RUN] Command not executed.")

    try:
        result = subprocess.run(
            command, shell=True, capture_output=True, text=True, timeout=30
        )
    except subprocess.TimeoutExpired:
        return error_response(
            "SHELL_TIMEOUT",
            "Command timed out after 30s",
            hint="Try a faster approach, increase timeout, or break work into smaller commands.",
        )

    max_lines = THRESH.get("shell_max_lines", 100)
    keep_head = THRESH.get("shell_keep_head", 30)
    keep_tail = THRESH.get("shell_keep_tail", 30)

    if result.returncode != 0:
        stderr = _strip_ansi(result.stderr)
        lines = stderr.strip().splitlines()
        if len(lines) > keep_tail:
            stderr = "\n".join(lines[-keep_tail:])
        return error_response(
            "SHELL_NONZERO",
            f"Exit code {result.returncode}",
            detail=stderr,
        )

    output = _strip_ansi(result.stdout or "")
    lines = output.strip().splitlines()
    truncated = len(lines) > max_lines

    if truncated:
        head = "\n".join(lines[:keep_head])
        tail = "\n".join(lines[-keep_tail:])
        output = f"{head}\n... [{len(lines) - keep_head - keep_tail} lines omitted] ...\n{tail}"

    return success_response(
        output, truncated=truncated, metadata={"returncode": result.returncode}
    )


@mcp.tool(name="lean_analyze_workspace")
def analyze_workspace(path: str = ".", max_depth: Optional[int] = None) -> str:
    """Lists files and folders under path (or returns metadata if path is a file)."""
    resolved = Path(path).resolve()
    if not resolved.exists():
        return error_response("PATH_NOT_FOUND", "Directory does not exist", str(resolved))

    if resolved.is_file():
        stat = resolved.stat()
        return success_response(
            {
                "type": "file",
                "name": resolved.name,
                "size_bytes": stat.st_size,
                "path": str(resolved),
            }
        )

    blacklist = {".git", "node_modules", "__pycache__", "venv", ".venv", "dist", "build", ".tox"}
    depth = max_depth if max_depth is not None else THRESH.get("workspace_max_depth", 10)
    max_files = THRESH.get("workspace_max_files", 150)

    files, dirs = [], []
    abort = False

    def walk(current: Path, current_depth: int):
        nonlocal abort
        if abort or current_depth > depth:
            return
        try:
            entries = sorted(current.iterdir(), key=lambda e: (e.is_file(), e.name.lower()))
        except PermissionError:
            return

        for entry in entries:
            if entry.name in blacklist or abort:
                continue
            rel_path = str(entry.relative_to(resolved))
            if entry.is_dir():
                dirs.append(rel_path)
                walk(entry, current_depth + 1)
            else:
                files.append(rel_path)
                if len(files) >= max_files:
                    abort = True
                    return

    walk(resolved, 0)
    return success_response(
        {"files": files, "directories": dirs},
        truncated=abort,
        metadata={"total_files": len(files), "total_directories": len(dirs), "root": str(resolved)},
    )


# ===========================================================================
# File Tools (Item 1: Split Dedicated Tools, Item 3: RAG, Item 4: Line Replacement)
# ===========================================================================
@mcp.tool(name="lean_file_read")
def file_read(
    path: str,
    start_line: int = 1,
    end_line: Optional[int] = None,
    query: Optional[str] = None,
) -> str:
    """Reads lines from a text file with line numbers. Supports query filtering (Item 3)."""
    resolved = Path(path).resolve()
    if not resolved.exists():
        return error_response("FILE_NOT_FOUND", "Path does not exist", str(resolved))

    try:
        lines = resolved.read_text(encoding="utf-8").splitlines()
        total_lines = len(lines)

        if query and query.strip():
            # Item 3: In-tool query filtering on numbered lines
            numbered_lines = [f"{idx}: {line}" for idx, line in enumerate(lines, start=1)]
            filtered, truncated = _filter_by_query(numbered_lines, query)
            return success_response(
                "\n".join(filtered),
                truncated=truncated,
                metadata={"total_lines": total_lines, "filtered_by_query": query},
            )

        start_line = max(1, start_line)
        page_size = THRESH.get("file_page_size", 200)
        window_end = end_line or (start_line + page_size - 1)
        actual_end = min(window_end, total_lines)

        page_lines = lines[start_line - 1 : actual_end]
        numbered_output = "\n".join(
            f"{idx}: {line}" for idx, line in enumerate(page_lines, start=start_line)
        )
        has_more = actual_end < total_lines

        return success_response(
            numbered_output,
            truncated=has_more,
            metadata={
                "start_line": start_line,
                "end_line": actual_end,
                "total_lines": total_lines,
                "has_more": has_more,
            },
        )
    except Exception as e:
        return error_response("FILE_READ_ERROR", f"Cannot read file: {e}", str(resolved))


@mcp.tool(name="lean_file_write")
def file_write(path: str, content: str, dry_run: bool = False) -> str:
    """Overwrites (or creates) a text file with the given content."""
    resolved = Path(path).resolve()
    if dry_run:
        return success_response({"path": str(resolved)}, message="[DRY RUN] Would write file.")
    resolved.parent.mkdir(parents=True, exist_ok=True)
    resolved.write_text(content, encoding="utf-8")
    return success_response(
        {"path": str(resolved), "words": len(content.split())},
        message="File written successfully.",
    )


@mcp.tool(name="lean_file_append")
def file_append(path: str, content: str, dry_run: bool = False) -> str:
    """Appends content to the end of an existing text file."""
    resolved = Path(path).resolve()
    if not resolved.exists():
        return error_response("FILE_NOT_FOUND", "Path does not exist", str(resolved))
    if dry_run:
        return success_response({"path": str(resolved)}, message="[DRY RUN] Would append to file.")
    with open(resolved, "a", encoding="utf-8") as f:
        f.write(content)
    return success_response(
        {"path": str(resolved), "appended_words": len(content.split())},
        message="Content appended successfully.",
    )


@mcp.tool(name="lean_file_replace_str")
def file_replace_str(path: str, old_str: str, new_str: str, dry_run: bool = False) -> str:
    """Replaces exact string occurrences in a file."""
    resolved = Path(path).resolve()
    if not resolved.exists():
        return error_response("FILE_NOT_FOUND", "Path does not exist", str(resolved))

    file_text = resolved.read_text(encoding="utf-8")
    occurrences = file_text.count(old_str)
    if occurrences == 0:
        return error_response(
            "TARGET_NOT_FOUND",
            "Target string 'old_str' was not found in the file.",
        )

    if dry_run:
        return success_response(
            {"occurrences": occurrences, "path": str(resolved)},
            message=f"[DRY RUN] Would replace {occurrences} occurrence(s).",
        )

    updated_text = file_text.replace(old_str, new_str)
    resolved.write_text(updated_text, encoding="utf-8")
    return success_response(
        {"path": str(resolved), "replacements_made": occurrences},
        message=f"Successfully replaced {occurrences} occurrence(s).",
    )


@mcp.tool(name="lean_file_replace_lines")
def file_replace_lines(
    path: str,
    start_line: int,
    end_line: int,
    new_content: str,
    dry_run: bool = False,
) -> str:
    """Item 4: Replaces a specific line range (1-indexed, inclusive) with new content."""
    resolved = Path(path).resolve()
    if not resolved.exists():
        return error_response("FILE_NOT_FOUND", "Path does not exist", str(resolved))

    lines = resolved.read_text(encoding="utf-8").splitlines()
    total_lines = len(lines)

    if start_line < 1 or start_line > total_lines or end_line < start_line:
        return error_response(
            "OUT_OF_BOUNDS",
            f"Invalid line range {start_line}-{end_line} for file with {total_lines} lines.",
        )

    actual_end = min(end_line, total_lines)

    if dry_run:
        return success_response(
            {"start_line": start_line, "end_line": actual_end, "total_lines": total_lines},
            message="[DRY RUN] Would replace specified line range.",
        )

    new_lines = new_content.splitlines() if new_content else []
    lines[start_line - 1 : actual_end] = new_lines

    resolved.write_text("\n".join(lines) + ("\n" if lines else ""), encoding="utf-8")
    return success_response(
        {
            "path": str(resolved),
            "replaced_range": f"{start_line}-{actual_end}",
            "lines_removed": actual_end - start_line + 1,
            "lines_added": len(new_lines),
            "new_total_lines": len(lines),
        },
        message="Line range replaced successfully.",
    )


# ===========================================================================
# Web & Search Tools
# ===========================================================================
@mcp.tool(name="lean_fallback_web_search")
def fallback_web_search(query: str) -> str:
    """Searches the web via DuckDuckGo and returns title/domain/url/snippet for the top results."""
    try:
        from ddgs import DDGS
        raw_results = list(DDGS().text(query, max_results=4))
    except Exception as e:
        return error_response("SEARCH_FAILED", "DuckDuckGo query failed", str(e))

    cleaned_results = []
    for r in raw_results:
        clean_url = re.sub(r"\?.*$", "", r.get("href", ""))
        domain = urlparse(clean_url).netloc
        cleaned_results.append(
            {
                "title": r.get("title", ""),
                "domain": domain,
                "url": clean_url,
                "snippet": r.get("body", ""),
            }
        )

    return success_response(cleaned_results, metadata={"query": query})


@mcp.tool(name="lean_web_scrape")
def web_scrape(
    url: str,
    query: Optional[str] = None,
    include_links: bool = False,
) -> str:
    """Scrapes clean article body. Supports optional query filtering (Item 3)."""
    if url.lower().endswith(".pdf"):
        return success_response(
            {"url": url},
            message="URL targets a PDF document. Please use pdf_read_text directly.",
        )

    try:
        import httpx
        response = httpx.get(
            url,
            timeout=30,
            follow_redirects=True,
            headers={"User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36"}
        )
        response.raise_for_status()
    except Exception as e:
        return error_response("SCRAPE_HTTP_ERROR", f"Failed to fetch URL: {e}", url)

    try:
        import trafilatura
        extracted = trafilatura.extract(
            response.text,
            output_format="markdown",
            favor_precision=True,
            include_links=include_links,
            include_images=False,
            include_tables=True,
        )
    except Exception:
        extracted = None

    if not extracted or not extracted.strip():
        return error_response("SCRAPE_EMPTY", "No extractable body content found.", url)

    if not include_links:
        extracted = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', extracted)

    paragraphs = [p.strip() for p in re.sub(r"\n{3,}", "\n\n", extracted.strip()).split("\n\n") if p.strip()]

    # Item 3: Query filtering on webpage paragraphs
    if query and query.strip():
        filtered_paras, truncated = _filter_by_query(paragraphs, query)
        return success_response(
            "\n\n".join(filtered_paras),
            truncated=truncated,
            metadata={"url": url, "filtered_by_query": query},
        )

    full_text = "\n\n".join(paragraphs)
    max_cap = THRESH.get("scrape_max_chars", 12000)
    truncated = len(full_text) > max_cap

    return success_response(
        full_text[:max_cap] if truncated else full_text,
        truncated=truncated,
        metadata={"url": url, "char_count": len(full_text)},
    )


# ===========================================================================
# Excel Manager Tools (Item 1: Split Dedicated Tools, Item 3: Query RAG)
# ===========================================================================
@mcp.tool(name="lean_excel_inspect")
def excel_inspect(path: str) -> str:
    """Returns sheet names, dimensions, and header row for an Excel workbook."""
    resolved = Path(path).resolve()
    if not resolved.exists():
        return error_response("XLSX_NOT_FOUND", "Spreadsheet does not exist", str(resolved))

    import openpyxl
    try:
        wb = openpyxl.load_workbook(resolved)
        active_sheet = wb.sheetnames[0]
        ws = wb[active_sheet]
        first_rows = list(ws.iter_rows(max_row=1, values_only=True))
        headers = list(first_rows[0]) if first_rows else []
        meta = {
            "sheet_names": wb.sheetnames,
            "active_sheet": active_sheet,
            "headers": headers,
            "dimensions": ws.dimensions,
            "max_rows": ws.max_row,
            "max_cols": ws.max_column,
        }
        wb.close()
        return success_response(meta)
    except Exception as e:
        return error_response("XLSX_CORRUPT", f"Cannot open workbook: {e}", str(resolved))


@mcp.tool(name="lean_excel_read_rows")
def excel_read_rows(
    path: str,
    sheet: Optional[str] = None,
    range_box: Optional[str] = None,
    output_format: Literal["json", "csv"] = "json",
    query: Optional[str] = None,
) -> str:
    """Reads rows from an Excel file as structured JSON. Supports query filtering (Item 3)."""
    resolved = Path(path).resolve()
    if not resolved.exists():
        return error_response("XLSX_NOT_FOUND", "Spreadsheet does not exist", str(resolved))

    import openpyxl
    try:
        wb = openpyxl.load_workbook(resolved)
    except Exception as e:
        return error_response("XLSX_CORRUPT", f"Cannot open workbook: {e}", str(resolved))

    ws_name = sheet or wb.sheetnames[0]
    if ws_name not in wb.sheetnames:
        wb.close()
        return error_response("XLSX_BAD_SHEET", f"Sheet '{ws_name}' not found", str(resolved))

    ws = wb[ws_name]
    iter_kwargs: Dict[str, Any] = {"values_only": True}
    if range_box:
        try:
            min_col, min_row, max_col, max_row = openpyxl.utils.cell.range_boundaries(range_box)
            iter_kwargs.update(
                min_row=min_row,
                max_row=min(max_row, ws.max_row or 1),
                min_col=min_col,
                max_col=min(max_col, ws.max_column or 1),
            )
        except Exception as e:
            wb.close()
            return error_response("XLSX_BAD_RANGE", f"Invalid range '{range_box}': {e}", str(resolved))

    raw_rows = list(ws.iter_rows(**iter_kwargs))
    wb.close()

    if not raw_rows:
        return success_response([])

    headers = [str(c) if c is not None else f"col_{i+1}" for i, c in enumerate(raw_rows[0])]
    dict_rows = []
    for row in raw_rows[1:]:
        dict_rows.append(
            {headers[i]: row[i] if i < len(row) else None for i in range(len(headers))}
        )

    # Item 3: Query filtering on dictionary rows
    if query and query.strip():
        row_strings = [json.dumps(r) for r in dict_rows]
        filtered_strs, truncated = _filter_by_query(row_strings, query)
        filtered_dicts = []
        for s in filtered_strs:
            if s.startswith("{"):
                try:
                    filtered_dicts.append(json.loads(s))
                except Exception:
                    pass
        return success_response(
            filtered_dicts or filtered_strs,
            truncated=truncated,
            metadata={"filtered_by_query": query},
        )

    if output_format == "json":
        return success_response(dict_rows)
    else:
        output = io.StringIO()
        writer = csv.writer(output)
        writer.writerow(headers)
        for r in dict_rows:
            writer.writerow([r.get(h, "") for h in headers])
        return success_response(output.getvalue())


@mcp.tool(name="lean_excel_write_rows")
def excel_write_rows(
    path: str,
    csv_data: str,
    sheet: Optional[str] = None,
    dry_run: bool = False,
) -> str:
    """Writes CSV-formatted rows into an Excel sheet, creating the workbook/sheet if needed."""
    resolved = Path(path).resolve()
    if not resolved.exists():
        return error_response("XLSX_NOT_FOUND", "Spreadsheet does not exist", str(resolved))

    import openpyxl
    try:
        wb = openpyxl.load_workbook(resolved)
    except Exception as e:
        return error_response("XLSX_CORRUPT", f"Cannot open workbook: {e}", str(resolved))

    ws_name = sheet or wb.sheetnames[0]
    if ws_name not in wb.sheetnames:
        wb.close()
        return error_response("XLSX_BAD_SHEET", f"Sheet '{ws_name}' not found", str(resolved))

    ws = wb[ws_name]
    reader = csv.reader(io.StringIO(csv_data.strip()))
    parsed_rows = list(reader)

    if dry_run:
        wb.close()
        return success_response(
            {"rows_to_write": len(parsed_rows)}, message="[DRY RUN] Would write rows."
        )

    for row in parsed_rows:
        ws.append(row)

    wb.save(resolved)
    wb.close()
    return success_response(
        {"rows_appended": len(parsed_rows), "path": str(resolved)},
        message="Successfully appended rows.",
    )


# ===========================================================================
# Word Manager Tools (Item 1: Split Dedicated Tools, Item 3: Query RAG)
# ===========================================================================
def _set_doc_accessibility_meta(doc: Any, title: Optional[str], language: str) -> None:
    if title:
        doc.core_properties.title = title
    try:
        styles_element = doc.styles.element
        lang_element = OxmlElement("w:lang")
        lang_element.set(qn("w:val"), language)
        styles_element.append(lang_element)
    except Exception:
        pass


def _make_table_accessible(table: Any) -> None:
    if not table.rows:
        return
    header_tr = table.rows[0]._tr.get_or_add_trPr()
    header_tr.append(OxmlElement("w:tblHeader"))
    for row in table.rows:
        trPr = row._tr.get_or_add_trPr()
        trPr.append(OxmlElement("w:cantSplit"))


def _add_markdown_paragraph(doc: Any, text: str, style: str = "Normal") -> None:
    p = doc.add_paragraph(style=style)
    tokens = re.split(r"(\*\*.*?\*\*|\*.*?\*)", text)
    for token in tokens:
        if token.startswith("**") and token.endswith("**"):
            run = p.add_run(token[2:-2])
            run.bold = True
        elif token.startswith("*") and token.endswith("*"):
            run = p.add_run(token[1:-1])
            run.italic = True
        else:
            p.add_run(token)


def _read_docx_robust(doc: Any) -> List[Dict[str, Any]]:
    extracted = []
    all_p_nodes = doc.element.body.xpath(".//w:p")

    for p_node in all_p_nodes:
        text = "".join(p_node.xpath(".//w:t/text()")).strip()
        if not text:
            continue

        heading_level = None
        style_val = p_node.xpath("./w:pPr/w:pStyle/@w:val")
        if style_val:
            style_name = str(style_val[0])
            if "Heading" in style_name or "heading" in style_name:
                digits = re.findall(r"\d+", style_name)
                if digits:
                    heading_level = int(digits[0])

        if heading_level is None:
            outline_lvl = p_node.xpath("./w:pPr/w:outlineLvl/@w:val")
            if outline_lvl:
                heading_level = int(outline_lvl[0]) + 1

        if heading_level is None:
            is_bold = bool(p_node.xpath(".//w:rPr/w:b"))
            font_sizes = p_node.xpath(".//w:rPr/w:sz/@w:val")
            max_size = max([int(s) for s in font_sizes], default=22)

            if is_bold and max_size >= 28:
                heading_level = 1
            elif is_bold and max_size >= 24:
                heading_level = 2
            elif is_bold and max_size >= 20:
                heading_level = 3
            elif is_bold and max_size >= 18:
                heading_level = 4

        extracted.append(
            {
                "type": "heading" if heading_level else "paragraph",
                "level": heading_level,
                "text": text,
            }
        )
    return extracted


@mcp.tool(name="lean_word_read_text")
def word_read_text(path: str, query: Optional[str] = None) -> str:
    """Reads body text from Word docx. Supports query filtering (Item 3)."""
    resolved = Path(path).resolve()
    if not resolved.exists():
        return error_response("DOCX_NOT_FOUND", "Document does not exist", str(resolved))

    try:
        doc = Document(str(resolved))
    except Exception as e:
        return error_response("DOCX_CORRUPT", f"Cannot open docx: {e}", str(resolved))

    standard_elements = []
    try:
        for p in doc.paragraphs:
            if p.text.strip():
                standard_elements.append(p.text.strip())
    except Exception:
        standard_elements = []

    if not standard_elements:
        robust_objs = _read_docx_robust(doc)
        paragraphs = [el["text"] for el in robust_objs]
        read_method = "robust_xml_fallback"
    else:
        paragraphs = standard_elements
        read_method = "standard"

    # Item 3: Query filtering on document paragraphs
    if query and query.strip():
        filtered, truncated = _filter_by_query(paragraphs, query)
        return success_response(
            filtered,
            truncated=truncated,
            metadata={"read_method": read_method, "filtered_by_query": query},
        )

    return success_response(
        paragraphs, metadata={"read_method": read_method, "total_paragraphs": len(paragraphs)}
    )


@mcp.tool(name="lean_word_read_outline")
def word_read_outline(path: str) -> str:
    """Returns the heading structure (outline) of a Word document."""
    resolved = Path(path).resolve()
    if not resolved.exists():
        return error_response("DOCX_NOT_FOUND", "Document does not exist", str(resolved))

    try:
        doc = Document(str(resolved))
    except Exception as e:
        return error_response("DOCX_CORRUPT", f"Cannot open docx: {e}", str(resolved))

    standard_elements = []
    try:
        for p in doc.paragraphs:
            if not p.text.strip():
                continue
            style_name = p.style.name if p.style else ""
            level = None
            if style_name.startswith("Heading"):
                try:
                    level = int(style_name.replace("Heading ", ""))
                except ValueError:
                    pass
            standard_elements.append({"type": "heading" if level else "paragraph", "level": level, "text": p.text})
    except Exception:
        standard_elements = []

    if not standard_elements:
        elements = _read_docx_robust(doc)
        read_method = "robust_xml_fallback"
    else:
        elements = standard_elements
        read_method = "standard"

    outline = [
        {"level": el["level"], "text": el["text"]}
        for el in elements
        if el["type"] == "heading" and el["level"] is not None and 1 <= el["level"] <= 4
    ]
    return success_response(outline, metadata={"read_method": read_method})


@mcp.tool(name="lean_word_write_doc")
def word_write_doc(
    path: str,
    doc_text: Optional[str] = None,
    write_mode: Literal["create", "append"] = "create",
    title: Optional[str] = None,
    language: str = "en-US",
) -> str:
    """Writes or appends to a Word document with WCAG 2.2 AA accessibility structures."""
    resolved = Path(path).resolve()
    if write_mode == "create":
        doc = Document()
        doc_title = title or resolved.stem
        doc.add_heading(doc_title, level=0)
        _set_doc_accessibility_meta(doc, title=doc_title, language=language)
    else:
        if not resolved.exists():
            return error_response("DOCX_NOT_FOUND", "Document does not exist", str(resolved))
        doc = Document(str(resolved))
        _set_doc_accessibility_meta(doc, title=title, language=language)

    if doc_text:
        lines = doc_text.strip().splitlines()
        i = 0
        while i < len(lines):
            line = lines[i].strip()
            if not line:
                i += 1
                continue

            if line.startswith("|") and line.endswith("|"):
                table_lines = []
                while i < len(lines) and lines[i].strip().startswith("|") and lines[i].strip().endswith("|"):
                    table_lines.append(lines[i].strip())
                    i += 1

                rows_data = []
                for tline in table_lines:
                    if re.match(r"^\|[\s\-:\t|]+\|$", tline):
                        continue
                    cells = [c.strip() for c in tline.split("|")[1:-1]]
                    rows_data.append(cells)

                if rows_data:
                    num_cols = max(len(r) for r in rows_data)
                    table = doc.add_table(rows=len(rows_data), cols=num_cols)
                    for r_idx, r_data in enumerate(rows_data):
                        for c_idx, c_val in enumerate(r_data):
                            if c_idx < num_cols:
                                table.cell(r_idx, c_idx).text = c_val
                    _make_table_accessible(table)
                continue

            elif line.startswith("#### "):
                doc.add_heading(line[5:], level=4)
            elif line.startswith("### "):
                doc.add_heading(line[4:], level=3)
            elif line.startswith("## "):
                doc.add_heading(line[3:], level=2)
            elif line.startswith("# "):
                doc.add_heading(line[2:], level=1)
            elif line.startswith("- ") or line.startswith("* "):
                _add_markdown_paragraph(doc, line[2:], style="List Bullet")
            elif re.match(r"^\d+\.\s", line):
                item_text = re.sub(r"^\d+\.\s", "", line)
                _add_markdown_paragraph(doc, item_text, style="List Number")
            else:
                _add_markdown_paragraph(doc, line, style="Normal")

            i += 1

    resolved.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(resolved))
    return success_response(
        {"path": str(resolved), "mode": write_mode, "language": language},
        message="Document saved with WCAG accessibility metadata.",
    )


# ===========================================================================
# PDF Manager Tools (Item 1: Split Dedicated Tools, Item 3: Query RAG)
# ===========================================================================
@mcp.tool(name="lean_pdf_read_text")
def pdf_read_text(
    path: str,
    start_page: int = 1,
    end_page: Optional[int] = None,
    query: Optional[str] = None,
) -> str:
    """Reads PDF text with PyMuPDF layout analysis. Supports page ranges and query filtering (Item 3)."""
    resolved = Path(path).resolve()
    if not resolved.exists():
        return error_response("PDF_NOT_FOUND", "PDF does not exist", str(resolved))

    try:
        doc = fitz.open(str(resolved))
    except Exception as e:
        return error_response("PDF_CORRUPT", f"Cannot parse PDF: {e}", str(resolved))

    if doc.is_encrypted:
        return error_response("PDF_ENCRYPTED", "PDF is password protected", str(resolved))

    total_pages = len(doc)
    s_page = max(1, start_page)
    e_page = min(total_pages, end_page) if end_page else total_pages

    extracted_pages = []
    for pno in range(s_page - 1, e_page):
        page = doc[pno]
        text = page.get_text("markdown") or page.get_text()
        extracted_pages.append(f"--- Page {pno + 1} ---\n" + text.strip())

    doc.close()

    # Item 3: Query filtering on PDF page text
    if query and query.strip():
        filtered, truncated = _filter_by_query(extracted_pages, query)
        return success_response(
            filtered,
            truncated=truncated,
            metadata={"start_page": s_page, "end_page": e_page, "filtered_by_query": query},
        )

    return success_response(
        extracted_pages,
        metadata={"start_page": s_page, "end_page": e_page, "total_pages": total_pages},
    )


@mcp.tool(name="lean_pdf_read_outline")
def pdf_read_outline(path: str) -> str:
    """Returns the table-of-contents/bookmark outline of a PDF, if it has one."""
    resolved = Path(path).resolve()
    if not resolved.exists():
        return error_response("PDF_NOT_FOUND", "PDF does not exist", str(resolved))

    try:
        doc = fitz.open(str(resolved))
    except Exception as e:
        return error_response("PDF_CORRUPT", f"Cannot parse PDF: {e}", str(resolved))

    if doc.is_encrypted:
        return error_response("PDF_ENCRYPTED", "PDF is password protected", str(resolved))

    toc = doc.get_toc()
    outline = [{"level": item[0], "title": item[1], "page": item[2]} for item in toc]
    doc.close()
    return success_response(outline)


# ===========================================================================
# Cache Manager Tools (Item 1: Split Dedicated Tools)
# ===========================================================================
@mcp.tool(name="lean_cache_list")
def cache_list() -> str:
    """Lists files currently stored in the scratch cache directory with their sizes."""
    if not CACHE_DIR.exists():
        return success_response([])
    entries = [
        {"filename": f.name, "size_bytes": f.stat().st_size}
        for f in sorted(CACHE_DIR.iterdir())
        if f.is_file()
    ]
    return success_response(entries)


@mcp.tool(name="lean_cache_view")
def cache_view(filename: str) -> str:
    """Reads the text content of a file previously stored in the scratch cache directory."""
    file_path = CACHE_DIR / filename
    if not file_path.exists():
        return error_response("CACHE_MISS", f"File '{filename}' not found.")
    text = file_path.read_text(encoding="utf-8")
    return success_response(text)


@mcp.tool(name="lean_cache_delete")
def cache_delete(filename: str) -> str:
    """Deletes a single file from the scratch cache directory."""
    file_path = CACHE_DIR / filename
    if file_path.exists():
        file_path.unlink()
        return success_response({"deleted": filename})
    return error_response("CACHE_MISS", f"File '{filename}' not found.")


@mcp.tool(name="lean_cache_clear")
def cache_clear() -> str:
    """Deletes every file in the scratch cache directory and returns how many were removed."""
    count = 0
    if CACHE_DIR.exists():
        for f in CACHE_DIR.iterdir():
            if f.is_file():
                f.unlink()
                count += 1
    return success_response({"files_removed": count})


# ===========================================================================
# Scratchpad Tools (Item 1: Split Dedicated Tools)
# ===========================================================================
@mcp.tool(name="lean_scratchpad_set")
def scratchpad_set(key: str, value: str) -> str:
    """Stores a key/value pair in the persistent scratchpad for recall across turns."""
    data = _load_scratchpad()
    data[key] = value
    _save_scratchpad(data)
    return success_response({"key": key}, message="Stored successfully.")


@mcp.tool(name="lean_scratchpad_get")
def scratchpad_get(key: str) -> str:
    """Retrieves a previously stored scratchpad value by key."""
    data = _load_scratchpad()
    if key not in data:
        return error_response("KEY_NOT_FOUND", f"Key '{key}' not in scratchpad.")
    return success_response({"key": key, "value": data[key]})


@mcp.tool(name="lean_scratchpad_delete")
def scratchpad_delete(key: str) -> str:
    """Deletes a key from the persistent scratchpad."""
    data = _load_scratchpad()
    if key not in data:
        return error_response("KEY_NOT_FOUND", f"Key '{key}' not in scratchpad.")
    del data[key]
    _save_scratchpad(data)
    return success_response({"deleted_key": key})


@mcp.tool(name="lean_scratchpad_list")
def scratchpad_list() -> str:
    """Lists all keys currently stored in the persistent scratchpad."""
    data = _load_scratchpad()
    return success_response(list(data.keys()))


# ===========================================================================
# Entry point
# ===========================================================================
def main() -> None:
    mcp.run()


if __name__ == "__main__":
    main()